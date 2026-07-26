"""Tests for the DOCX extractor and the audio-transcription guards."""

from __future__ import annotations

from pathlib import Path

import pytest

from llm_wiki.parsers import audio
from llm_wiki.parsers.audio import TranscriptionError, transcribe_audio
from llm_wiki.parsers.docx import ParseError, parse_docx


def _make_docx(
    path: Path, paragraphs: list[str], table: list[list[str]] | None = None
) -> None:
    from docx import Document

    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    if table:
        tbl = doc.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, value in enumerate(row):
                tbl.rows[r].cells[c].text = value
    doc.save(str(path))


# ── DOCX ────────────────────────────────────────────────────────────────────


def test_parse_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "doc.docx"
    _make_docx(
        path,
        ["Первый абзац документа.", "Второй абзац."],
        table=[["Метрика", "Значение"], ["Выручка", "100"]],
    )
    text = parse_docx(path)
    assert "Первый абзац документа." in text
    assert "Второй абзац." in text
    # table rows are flattened to "a | b"
    assert "Метрика | Значение" in text
    assert "Выручка | 100" in text


def test_parse_docx_empty_raises(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    _make_docx(path, [])
    with pytest.raises(ParseError):
        parse_docx(path)


def test_parse_docx_not_a_docx_raises(tmp_path: Path) -> None:
    path = tmp_path / "fake.docx"
    path.write_text("this is plainly not a zip/docx", encoding="utf-8")
    with pytest.raises(ParseError):
        parse_docx(path)


# ── Audio guards (no network) ─────────────────────────────────────────────────


def test_transcribe_missing_key_raises(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(audio.settings, "openai_api_key", "")
    path = tmp_path / "clip.mp3"
    path.write_bytes(b"\x00\x01")
    with pytest.raises(TranscriptionError, match="OPENAI_API_KEY"):
        transcribe_audio(path)


def test_transcribe_rejects_oversized(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(audio.settings, "openai_api_key", "sk-test")
    path = tmp_path / "big.mp3"
    # Sparse file just over the 25 MB API limit — no real bytes written.
    with path.open("wb") as handle:
        handle.seek(audio._MAX_AUDIO_BYTES + 1)
        handle.write(b"\x00")
    with pytest.raises(TranscriptionError, match="25 MB"):
        transcribe_audio(path)
