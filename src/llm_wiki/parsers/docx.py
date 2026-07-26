"""DOCX (Word) text extractor via python-docx.

Pulls paragraph text plus table-cell text — business documents often keep key
facts in tables, so dropping them silently would lose information. The extracted
plain text feeds the same ingestion pipeline as PDF/Markdown; the Writer Agent
restructures it into a wiki page, so we do not need to preserve exact layout.
"""

from __future__ import annotations

from pathlib import Path

import structlog

logger = structlog.get_logger(__name__)


class ParseError(Exception):
    """Raised when the DOCX cannot be opened or yields no usable text."""


def parse_docx(path: Path) -> str:
    """Extract plain text (paragraphs + tables) from a .docx file.

    Args:
        path: Absolute path to the Word document.

    Returns:
        Extracted plain text (paragraphs first, then table rows as ``a | b | c``).

    Raises:
        ParseError: If python-docx is missing, the file cannot be opened, or no
            text could be extracted.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency always installed
        raise ParseError("python-docx is not installed") from exc

    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001 - surface any python-docx/zip error
        raise ParseError(f"Could not open DOCX {path}: {exc}") from exc

    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                blocks.append(line)

    text = "\n\n".join(blocks).strip()
    if not text:
        raise ParseError(f"No extractable text in DOCX {path}")
    logger.info("docx_parsed", path=str(path), chars=len(text))
    return text
